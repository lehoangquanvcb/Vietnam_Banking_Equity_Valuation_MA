
from pathlib import Path
from io import BytesIO
import os, math, json, tempfile
from xml.sax.saxutils import escape
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import FuncFormatter

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, Table, TableStyle, SimpleDocTemplate, Image as RLImage, Spacer, KeepTogether, PageBreak

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

try:
    from scripts.history_engine import load_effective_history, target_and_peer, coverage_note, BENCHMARK_LABEL
except Exception:
    from history_engine import load_effective_history, target_and_peer, coverage_note, BENCHMARK_LABEL

try:
    from scripts.narrative_engine import (
        executive_summary,business_profile,profitability_text,asset_quality_text,
        funding_text,capital_text,valuation_text,catalysts_risks,mna_text,pct,mult,money,bn,n
    )
except Exception:
    from narrative_engine import (
        executive_summary,business_profile,profitability_text,asset_quality_text,
        funding_text,capital_text,valuation_text,catalysts_risks,mna_text,pct,mult,money,bn,n
    )
try:
    from scripts.quality_engine import assess_report_quality, normalization_flags
except Exception:
    from quality_engine import assess_report_quality, normalization_flags
try:
    from scripts.strategic_case import load_research, strategic_reasonableness, reasonableness_conclusion, all_bank_means
except Exception:
    from strategic_case import load_research, strategic_reasonableness, reasonableness_conclusion, all_bank_means

PAGE_W,PAGE_H=A4
MARGIN=16*mm

# Chuẩn trình bày báo cáo: toàn bộ biểu đồ dùng Lato, cỡ 10.
plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Lato", "DejaVu Sans"],
    "font.size": 10,
    "axes.titlesize": 10,
    "axes.labelsize": 10,
    "xtick.labelsize": 10,
    "ytick.labelsize": 10,
    "legend.fontsize": 10,
})

METRIC_VI = {
    "TotalAssets":"Tổng tài sản",
    "GrossLoans":"Dư nợ khách hàng",
    "CustomerDeposits":"Tiền gửi khách hàng",
    "Equity":"Vốn chủ sở hữu",
    "TangibleEquity":"Vốn chủ sở hữu hữu hình",
    "NPAT":"Lợi nhuận sau thuế",
    "NetInterestIncome":"Thu nhập lãi thuần",
    "OperatingIncome":"Tổng thu nhập hoạt động",
    "ProvisionExpense":"Chi phí dự phòng",
    "ROE":"ROE","ROA":"ROA","NIM":"NIM",
    "NPL":"Tỷ lệ nợ xấu (NPL)",
    "CAR":"Hệ số an toàn vốn (CAR)",
    "CIR":"Tỷ lệ chi phí/thu nhập (CIR)",
    "LDR":"Tỷ lệ cho vay/tiền gửi (LDR)",
    "CASA":"Tỷ lệ tiền gửi không kỳ hạn (CASA)",
    "PB":"P/B","PE":"P/E",
}
METHOD_VI={
    "ResidualIncome":"Thu nhập thặng dư",
    "JustifiedPB":"P/B hợp lý",
    "PeerPB":"P/B nhóm so sánh",
    "HistoricalPB":"P/B lịch sử",
}
def _metric_vi(metric):
    return METRIC_VI.get(str(metric),str(metric))

def _vi_num(x,d=0):
    try:
        x=float(x)
        if not math.isfinite(x): return ""
    except Exception:
        return ""
    s=f"{x:,.{d}f}"
    return s.replace(",","§").replace(".",",").replace("§",".")

def _fmt_pct_axis(v,pos=None):
    return f"{_vi_num(v*100,1)}%"

def _fmt_vnd_axis(v,pos=None):
    return _vi_num(v,0)

def _fmt_decimal_axis(v,pos=None):
    av=abs(v)
    d=0 if av>=100 else 1 if av>=10 else 2
    return _vi_num(v,d)

def _set_y_padding(ax, values, percent=False):
    vals=pd.to_numeric(pd.Series(values),errors="coerce").dropna()
    if vals.empty:return
    lo=float(vals.min()); hi=float(vals.max())
    span=hi-lo
    if span<=0:
        span=max(abs(hi)*.20,.01 if percent else 1.0)
    pad=max(span*.15,.003 if percent else abs(hi)*.03 if hi else 1.0)
    lower=lo-pad; upper=hi+pad
    if percent and lo>=0: lower=max(0.0,lower)
    if lower==upper: upper=lower+(0.01 if percent else 1.0)
    ax.set_ylim(lower,upper)


def _load(root):
    root=Path(root); data=root/"data"; out=data/"model_outputs"
    def csv(p):
        try:return pd.read_csv(p)
        except Exception:return pd.DataFrame()
    return {
        "summary":csv(out/"valuation_summary.csv"),
        "methods":csv(out/"valuation_methods.csv"),
        "peer":csv(out/"peer_summary.csv"),
        "mna":csv(out/"mna_baseline.csv"),
        "hist":load_effective_history(root,csv(data/"bank_history_long.csv")),
        "prices":csv(data/"price_history.csv"),
        "precedents":csv(root/"config/transaction_precedents.csv"),
        "research":load_research(root),
        "cfg":json.loads((root/"config/model_config.json").read_text(encoding="utf-8")) if (root/"config/model_config.json").exists() else {}
    }

def _font_paths():
    """Ưu tiên Lato; nếu runtime chưa có Lato thì dùng DejaVu Sans Unicode để report vẫn xuất được."""
    lato_candidates = [
        ("/usr/share/fonts/truetype/lato/Lato-Regular.ttf","/usr/share/fonts/truetype/lato/Lato-Bold.ttf"),
        ("/usr/share/fonts/lato/Lato-Regular.ttf","/usr/share/fonts/lato/Lato-Bold.ttf"),
        (str(Path.home()/".fonts/Lato-Regular.ttf"),str(Path.home()/".fonts/Lato-Bold.ttf")),
        ("C:/Windows/Fonts/Lato-Regular.ttf","C:/Windows/Fonts/Lato-Bold.ttf"),
        ("C:/Windows/Fonts/Lato.ttf","C:/Windows/Fonts/Lato-Bold.ttf"),
    ]
    try:
        from matplotlib import font_manager
        lato_candidates.insert(0,(
            font_manager.findfont("Lato", fallback_to_default=False),
            font_manager.findfont(font_manager.FontProperties(family="Lato", weight="bold"), fallback_to_default=False)
        ))
    except Exception:
        pass
    for a,b in lato_candidates:
        if a and b and Path(a).exists() and Path(b).exists():
            return a,b,"Lato"

    fallback=[]
    try:
        import matplotlib
        mf=Path(matplotlib.get_data_path())/"fonts"/"ttf"
        fallback.append((str(mf/"DejaVuSans.ttf"),str(mf/"DejaVuSans-Bold.ttf")))
    except Exception:
        pass
    fallback += [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf","/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf","/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    ]
    for a,b in fallback:
        if a and b and Path(a).exists() and Path(b).exists():
            return a,b,"DejaVu Sans"
    return None,None,None

def report_font_status():
    reg,bold,name=_font_paths()
    return {"ok":bool(reg and bold),"font":name or "Không xác định","is_lato":name=="Lato"}

def _register_pdf_fonts():
    reg,bold,name=_font_paths()
    if not reg or not bold:
        raise RuntimeError("Không tìm thấy font Unicode để xuất PDF.")
    if "VNFont" not in pdfmetrics.getRegisteredFontNames(): pdfmetrics.registerFont(TTFont("VNFont",reg))
    if "VNFont-Bold" not in pdfmetrics.getRegisteredFontNames(): pdfmetrics.registerFont(TTFont("VNFont-Bold",bold))
    return "VNFont","VNFont-Bold"

def _period_sort(s):
    import re
    st=str(s).upper().strip()
    y=re.search(r"(20\d{2})",st)
    q=re.search(r"Q\s*([1-4])",st)
    return (int(y.group(1)) if y else 0,int(q.group(1)) if q else 0,st)

def _period_to_date(s):
    """Convert quarter labels such as 2025-Q2 / 2025Q2 to a true timestamp.
    Using a real date axis prevents categorical-axis reordering when target and peer
    series have different missing quarters.
    """
    y,q,_=_period_sort(s)
    if y and q:
        return pd.Timestamp(year=y,month=(q-1)*3+1,day=1)
    if y:
        return pd.Timestamp(year=y,month=1,day=1)
    return pd.NaT

def _with_period_date(df):
    if df is None or df.empty:
        return df
    out=df.copy()
    out["_date"]=out["Period"].map(_period_to_date)
    out=out.dropna(subset=["_date"]).sort_values(["_date","Metric"] if "Metric" in out.columns else ["_date"])
    return out

def _set_quarter_ticks(ax, frames, max_ticks=10):
    dates=[]
    for f in frames:
        if f is not None and len(f) and "_date" in f.columns:
            dates.extend(pd.to_datetime(f["_date"],errors="coerce").dropna().tolist())
    if not dates:
        return
    uniq=sorted(pd.unique(pd.Series(dates)))
    step=max(1,int(np.ceil(len(uniq)/max_ticks)))
    ticks=uniq[::step]
    if uniq[-1] not in ticks: ticks.append(uniq[-1])
    ax.set_xticks(ticks)
    ax.set_xticklabels([f"{pd.Timestamp(d).year}-Q{((pd.Timestamp(d).month-1)//3)+1}" for d in ticks],rotation=45,ha="right")

def _metric_history(hist,ticker,metrics):
    if hist.empty:return pd.DataFrame()
    x=hist[hist["Ticker"].astype(str).eq(str(ticker)) & hist["Metric"].astype(str).isin(metrics)].copy()
    if x.empty:return x
    x["Value"]=pd.to_numeric(x["Value"],errors="coerce")
    x=x.dropna(subset=["Value"])
    ratio_metrics={"ROE","ROA","NIM","NPL","CAR","CIR","LDR","CASA","PB","PE"}
    x=x[~(x["Metric"].astype(str).isin(ratio_metrics) & (x["Value"]==0))]
    x=x[~((x["Metric"].astype(str)=="CAR") & (x["Value"]<=0))]
    return _with_period_date(x)

def _peer_metric_history(hist,metrics):
    if hist.empty:return pd.DataFrame()
    x=hist[hist["Metric"].astype(str).isin(metrics)].copy()
    if x.empty:return x
    x["Value"]=pd.to_numeric(x["Value"],errors="coerce")
    x=x.dropna(subset=["Value"])
    ratio_metrics={"ROE","ROA","NIM","NPL","CAR","CIR","LDR","CASA","PB","PE"}
    x=x[~(x["Metric"].astype(str).isin(ratio_metrics) & (x["Value"]==0))]
    x=x[~((x["Metric"].astype(str)=="CAR") & (x["Value"]<=0))]
    x["_date"]=x["Period"].map(_period_to_date)
    x=x.dropna(subset=["_date"])
    out=x.groupby(["_date","Metric"],as_index=False).agg(Value=("Value","mean"),BankCount=("Ticker","nunique"))
    return out.sort_values(["_date","Metric"])

def _fig_to_png(fig):
    bio=BytesIO()
    fig.savefig(bio,format="png",dpi=160,bbox_inches="tight")
    plt.close(fig); bio.seek(0)
    return bio

def _chart_history(hist,ticker,metrics,title,percent=False):
    fig,ax=plt.subplots(figsize=(8.2,3.05)); plotted=[]; notes=[]
    balance_metrics={'TotalAssets','GrossLoans','CustomerDeposits','Equity','TangibleEquity'}
    use_trillion=(not percent) and all(str(m) in balance_metrics for m in metrics)
    scale=1e12 if use_trillion else 1.0
    any_target=False
    for metric in metrics:
        g,pg=target_and_peer(hist,ticker,metric)
        label=_metric_vi(metric)
        if len(g):
            any_target=True; gy=pd.to_numeric(g.Value,errors='coerce')/scale; plotted.extend(gy.dropna().tolist())
            if len(g)<3:
                line=ax.plot(g.PeriodDate,gy,linestyle='None',marker='o',markersize=5,label=f'{ticker} - {label}')[0]
            else:
                line=ax.plot(g.PeriodDate,gy,marker='o',markersize=3.8,linewidth=1.9,label=f'{ticker} - {label}')[0]
            if len(pg):
                py=pd.to_numeric(pg.PeerMean,errors='coerce')/scale; plotted.extend(py.dropna().tolist())
                if len(pg)<3:
                    ax.plot(pg.PeriodDate,py,linestyle='None',marker='x',markersize=5,color=line.get_color(),alpha=.72,label=f'{BENCHMARK_LABEL} - {label}')
                else:
                    ax.plot(pg.PeriodDate,py,linestyle='--',marker='o',markersize=3,linewidth=1.8,color=line.get_color(),alpha=.72,label=f'{BENCHMARK_LABEL} - {label}')
            notes.append(coverage_note(g,pg))
    if not any_target:
        ax.text(.5,.5,'Chưa có đủ dữ liệu lịch sử',ha='center',va='center');ax.set_axis_off()
    else:
        ax.set_title(title+' - so với '+BENCHMARK_LABEL);ax.grid(alpha=.20);ax.legend(ncol=2,loc='best')
        if percent:ax.yaxis.set_major_formatter(FuncFormatter(_fmt_pct_axis))
        elif use_trillion:ax.set_ylabel('Nghìn tỷ đồng');ax.yaxis.set_major_formatter(FuncFormatter(_fmt_decimal_axis))
        else:ax.yaxis.set_major_formatter(FuncFormatter(_fmt_decimal_axis))
        _set_y_padding(ax,plotted,percent=percent)
        ax.text(0,-0.30,' '.join(notes),transform=ax.transAxes,fontsize=8,color='#666666',va='top',wrap=True)
        fig.autofmt_xdate();fig.tight_layout(rect=[0,.08,1,1])
    return _fig_to_png(fig)


def _chart_peer(summary,ticker):
    fig,ax=plt.subplots(figsize=(8.2,2.85))
    q=summary.dropna(subset=["ROE_Used","PB_Current"]).copy()
    if q.empty:
        ax.text(.5,.5,"Chưa có đủ dữ liệu peer",ha="center",va="center"); ax.set_axis_off()
    else:
        ax.scatter(q["ROE_Used"],q["PB_Current"],s=45,alpha=.75)
        for _,r in q.iterrows():
            ax.annotate(str(r["Ticker"]),(r["ROE_Used"],r["PB_Current"]),xytext=(3,3),textcoords="offset points")
        cur=q[q["Ticker"].astype(str).eq(str(ticker))]
        if len(cur): ax.scatter(cur["ROE_Used"],cur["PB_Current"],s=120,marker="*")
        mean_roe=pd.to_numeric(q["ROE_Used"],errors="coerce").mean(); mean_pb=pd.to_numeric(q["PB_Current"],errors="coerce").mean()
        ax.axvline(mean_roe,linestyle="--",alpha=.65,label=f"ROE bình quân 20 ngân hàng niêm yết {_vi_num(mean_roe*100,1)}%"); ax.axhline(mean_pb,linestyle="--",alpha=.65,label=f"P/B bình quân 20 ngân hàng niêm yết {_vi_num(mean_pb,2)}x")
        ax.set_xlabel("ROE"); ax.set_ylabel("P/B (x)"); ax.set_title("Bản đồ ROE - P/B - so với bình quân 20 ngân hàng niêm yết"); ax.legend(fontsize=7)
        ax.xaxis.set_major_formatter(FuncFormatter(_fmt_pct_axis)); ax.yaxis.set_major_formatter(FuncFormatter(lambda v,pos:_vi_num(v,2))); ax.grid(alpha=.2)
    return _fig_to_png(fig)

def _chart_valuation(methods,row,summary=None):
    fig,ax=plt.subplots(figsize=(8.2,2.75))
    m=methods[methods["Ticker"].astype(str).eq(str(row.get("Ticker")))].copy() if len(methods) else pd.DataFrame()
    if m.empty:
        ax.text(.5,.5,"Chưa có kết quả theo phương pháp",ha="center",va="center"); ax.set_axis_off()
    else:
        m["FairValuePerShare"]=pd.to_numeric(m["FairValuePerShare"],errors="coerce"); m=m.dropna(subset=["FairValuePerShare"]); m["FairValuePerShare"]=m["FairValuePerShare"]*1000
        m["Phương pháp"]=m["Method"].map(lambda x: METHOD_VI.get(str(x),str(x)))
        ax.bar(m["Phương pháp"],m["FairValuePerShare"])
        if n(row.get("Price")) is not None: ax.axhline(n(row.get("Price"))*1000,linestyle="--",label="Giá thị trường")
        if summary is not None and len(summary) and n(row.get("BVPS_Used")) is not None:
            mpb=pd.to_numeric(summary["PB_Current"],errors="coerce").mean()
            if pd.notna(mpb): ax.axhline(mpb*n(row.get("BVPS_Used"))*1000,linestyle=":",linewidth=2,label=f"Giá hàm ý P/B bình quân 20 ngân hàng niêm yết ({_vi_num(mpb,2)}x)")
        slo=n(row.get("StrategicPriceLow")); shi=n(row.get("StrategicPriceHigh"))
        if slo is not None and shi is not None: ax.axhspan(slo*1000,shi*1000,alpha=.10,label="Vùng giá chiến lược/M&A")
        ax.set_title("Giá trị theo phương pháp - so với bình quân 20 ngân hàng niêm yết"); ax.set_ylabel("Đồng/cp"); ax.yaxis.set_major_formatter(FuncFormatter(_fmt_vnd_axis)); ax.tick_params(axis="x",rotation=20,labelsize=10); ax.legend(fontsize=7)
    return _fig_to_png(fig)


def _chart_scores(row,summary=None):
    names=["Sinh lời","Tăng trưởng","Chất lượng TS","Nguồn vốn","An toàn vốn","Định giá"]
    score_cols=["ProfitabilityScore","GrowthScore","AssetQualityScore","FundingScore","CapitalScore","ValuationScore"]
    vals=[row.get(c) for c in score_cols]; fig,ax=plt.subplots(figsize=(8.2,2.70)); y=np.arange(len(names)); vv=[n(v) or 0 for v in vals]
    ax.barh(y,vv); ax.set_yticks(y,names); ax.set_xlim(0,100); ax.set_title("Thẻ điểm 6 trụ cột · so với bình quân 20 ngân hàng niêm yết")
    for i,v in enumerate(vv): ax.text(v+1,i,f"{v:.0f}",va="center")
    if summary is not None and len(summary):
        means=[pd.to_numeric(summary[c],errors="coerce").mean() for c in score_cols]
        ax.plot(means,y,linestyle="--",marker="o",label="Bình quân 20 ngân hàng niêm yết"); ax.legend(fontsize=7)
    return _fig_to_png(fig)


def _chart_sensitivity(row):
    bv=n(row.get("BVPS_Used")); price=n(row.get("Price"))
    coe=n(row.get("COE")) or .13; g=n(row.get("LTG")) or .05; roe=n(row.get("NormalizedROE_Used")) or .15
    fig,ax=plt.subplots(figsize=(8.2,2.75))
    if bv is None:
        ax.text(.5,.5,"Chưa có BVPS để lập sensitivity",ha="center",va="center"); ax.set_axis_off()
    else:
        coes=np.array([coe-.02,coe-.01,coe,coe+.01,coe+.02])
        roes=np.array([roe-.04,roe-.02,roe,roe+.02,roe+.04])
        mat=np.zeros((len(roes),len(coes)))
        for i,r in enumerate(roes):
            for j,c in enumerate(coes):
                den=max(c-g,.01)
                pb=max(.3,min(4.0,(r-g)/den))
                mat[i,j]=pb*bv
        display_mat=mat*1000.0
        im=ax.imshow(display_mat,aspect="auto")
        ax.set_xticks(range(len(coes)),[f"{_vi_num(x*100,1)}%" for x in coes])
        ax.set_yticks(range(len(roes)),[f"{_vi_num(x*100,1)}%" for x in roes])
        ax.set_xlabel("Chi phí vốn chủ sở hữu (COE)"); ax.set_ylabel("ROE chuẩn hóa")
        ax.set_title("Sensitivity giá trị cơ bản: ROE x COE")
        for i in range(display_mat.shape[0]):
            for j in range(display_mat.shape[1]): ax.text(j,i,f"{_vi_num(display_mat[i,j]/1000,1)}k",ha="center",va="center")
        cb=fig.colorbar(im,ax=ax,fraction=.035,pad=.03,label="Đồng/cp"); cb.ax.yaxis.set_major_formatter(FuncFormatter(_fmt_vnd_axis))
    return _fig_to_png(fig)

def _chart_price(prices,ticker):
    fig,ax=plt.subplots(figsize=(8.2,2.75)); p=prices[prices["Ticker"].astype(str).eq(str(ticker))].copy() if len(prices) else pd.DataFrame()
    if p.empty:
        ax.text(.5,.5,"Chưa có lịch sử giá",ha="center",va="center"); ax.set_axis_off()
    else:
        allp=prices.copy(); allp["Date"]=pd.to_datetime(allp["Date"],errors="coerce"); allp["Close"]=pd.to_numeric(allp["Close"],errors="coerce"); allp=allp.dropna(subset=["Date","Close"]).sort_values(["Ticker","Date"])
        allp["Norm"]=allp.groupby("Ticker")["Close"].transform(lambda s: s/s.iloc[0]*100 if len(s) and s.iloc[0] else np.nan)
        bench=allp.groupby("Date",as_index=False)["Norm"].mean(); p=allp[allp["Ticker"].astype(str).eq(str(ticker))].copy()
        ax.plot(p["Date"],p["Norm"],label=f"{ticker} (chỉ số=100)"); ax.plot(bench["Date"],bench["Norm"],linestyle="--",label="Bình quân 20 ngân hàng niêm yết (chỉ số=100)")
        ax.set_title("Diễn biến giá tương đối - so với bình quân 20 ngân hàng niêm yết"); ax.set_ylabel("Chỉ số giá (đầu kỳ=100)"); ax.grid(alpha=.2); ax.legend(fontsize=7)
    return _fig_to_png(fig)


def _pdf_styles(font,bold):
    styles=getSampleStyleSheet()
    return {
        "h1":ParagraphStyle("h1",fontName=bold,fontSize=14,leading=18,textColor=colors.HexColor("#0F2747"),spaceAfter=5),
        "h2":ParagraphStyle("h2",fontName=bold,fontSize=12,leading=15,textColor=colors.HexColor("#0F2747"),spaceAfter=4),
        "body":ParagraphStyle("body",fontName=font,fontSize=11,leading=14.3,textColor=colors.HexColor("#222222"),alignment=TA_JUSTIFY,spaceBefore=6,spaceAfter=0),
        "small":ParagraphStyle("small",fontName=font,fontSize=10,leading=13,textColor=colors.HexColor("#555555"),alignment=TA_JUSTIFY),
        "kpi":ParagraphStyle("kpi",fontName=bold,fontSize=10,leading=13,textColor=colors.HexColor("#0F2747"),alignment=TA_CENTER),
    }

def _draw_para(c,text,style,x,y,w,h):
    safe_text=escape(str(text)).replace("\n","<br/>")
    p=Paragraph(safe_text,style)
    _,ph=p.wrap(w,h); p.drawOn(c,x,y-ph); return y-ph

def _draw_header(c,title,ticker,page,font,bold,stamp=None):
    c.setFillColor(colors.HexColor("#0F2747")); c.rect(0,PAGE_H-20*mm,PAGE_W,20*mm,fill=1,stroke=0)
    c.setFillColor(colors.white); c.setFont(bold,11); c.drawString(MARGIN,PAGE_H-12.5*mm,title)
    c.setFont(font,7.5); c.drawRightString(PAGE_W-MARGIN,PAGE_H-12.5*mm,f"{ticker} | Trang {page}/10")
    if stamp and stamp!="ĐỦ ĐIỀU KIỆN PHÁT HÀNH":
        c.saveState()
        c.setFillColor(colors.HexColor("#B42318")); c.setFont(bold,18)
        c.translate(PAGE_W/2,PAGE_H/2); c.rotate(32)
        c.drawCentredString(0,0,stamp)
        c.restoreState()
    c.setFillColor(colors.HexColor("#555555")); c.setFont(font,6.8)
    c.setFont(font,6.2)
    c.drawString(MARGIN,7*mm,"Nguồn: Vnstock/BCTC + mô hình nội bộ. Dữ liệu thực / số liệu tính toán / giả định được tách riêng.")
    c.drawRightString(PAGE_W-MARGIN,7*mm,"Tham khảo - không phải khuyến nghị đầu tư/chào mua.")

def _draw_image(c,bio,x,y,w,h):
    from reportlab.lib.utils import ImageReader
    bio.seek(0); c.drawImage(ImageReader(bio),x,y,width=w,height=h,preserveAspectRatio=True,anchor="c")

def _kpi_table(c,items,styles,y):
    data=[[Paragraph(k,styles["small"]),Paragraph(v,styles["kpi"])] for k,v in items]
    t=Table(data,colWidths=[31*mm,31*mm],rowHeights=12*mm)
    t.setStyle(TableStyle([
        ("GRID",(0,0),(-1,-1),.35,colors.HexColor("#D9E2F3")),
        ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#F3F6FA")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("ALIGN",(1,0),(1,-1),"CENTER"),
    ]))
    tw,th=t.wrap(62*mm,80*mm); t.drawOn(c,PAGE_W-MARGIN-tw,y-th); return y-th

def _report_context(root,ticker):
    d=_load(root); summary=d["summary"]
    row=summary[summary["Ticker"].astype(str).eq(str(ticker))]
    if row.empty: raise ValueError(f"Không có valuation_summary cho {ticker}")
    row=row.iloc[0].to_dict()
    peer_row=None
    if len(d["peer"]):
        q=d["peer"][d["peer"]["PeerGroup"].astype(str).eq(str(row.get("PeerGroup")))]
        if len(q): peer_row=q.iloc[0].to_dict()
    mna_row=None
    if len(d["mna"]):
        q=d["mna"][d["mna"]["Ticker"].astype(str).eq(str(ticker))]
        if len(q): mna_row=q.iloc[0].to_dict()
    return d,row,peer_row,mna_row

def _rl_flow_image(bio,max_width=178*mm):
    """Tạo ảnh ReportLab giữ đúng tỷ lệ và dùng hết chiều ngang vùng in."""
    from PIL import Image as PILImage
    bio.seek(0)
    raw=bio.getvalue()
    tmp=BytesIO(raw)
    im=PILImage.open(tmp)
    iw,ih=im.size
    ratio=ih/iw if iw else .4
    h=max_width*ratio
    bio.seek(0)
    return RLImage(bio,width=max_width,height=h)

def _pdf_page_decor(canvas_obj,doc_obj,ticker,font,bold,stamp=None):
    canvas_obj.saveState()
    canvas_obj.setFillColor(colors.HexColor("#0F2747"))
    canvas_obj.rect(0,PAGE_H-16*mm,PAGE_W,16*mm,fill=1,stroke=0)
    canvas_obj.setFillColor(colors.white)
    canvas_obj.setFont(bold,8.8)
    canvas_obj.drawString(MARGIN,PAGE_H-10.5*mm,"BÁO CÁO PHÂN TÍCH, ĐỊNH GIÁ & M&A NGÂN HÀNG")
    canvas_obj.setFont(font,7.4)
    canvas_obj.drawRightString(PAGE_W-MARGIN,PAGE_H-10.5*mm,f"{ticker} | Trang {doc_obj.page}")
    canvas_obj.setFillColor(colors.HexColor("#666666"))
    canvas_obj.setFont(font,5.8)
    canvas_obj.drawString(MARGIN,7*mm,"Nguồn: Vnstock/BCTC + mô hình nội bộ. Dữ liệu thực, số liệu tính toán và giả định được tách riêng.")
    canvas_obj.drawRightString(PAGE_W-MARGIN,7*mm,"Tham khảo - không phải khuyến nghị đầu tư/chào mua.")
    if stamp and stamp!="ĐỦ ĐIỀU KIỆN PHÁT HÀNH":
        canvas_obj.setFillColor(colors.HexColor("#B42318"))
        canvas_obj.setFont(bold,16)
        canvas_obj.translate(PAGE_W/2,PAGE_H/2); canvas_obj.rotate(30)
        canvas_obj.drawCentredString(0,0,stamp)
    canvas_obj.restoreState()

def _flow_table(data,font,bold,col_widths=None,font_size=7.5,header=True):
    t=Table(data,colWidths=col_widths,repeatRows=1 if header else 0,hAlign="LEFT")
    style=[
        ("GRID",(0,0),(-1,-1),.35,colors.HexColor("#CBD5E1")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("FONTNAME",(0,0),(-1,-1),font),
        ("FONTSIZE",(0,0),(-1,-1),font_size),
        ("LEFTPADDING",(0,0),(-1,-1),4),
        ("RIGHTPADDING",(0,0),(-1,-1),4),
        ("TOPPADDING",(0,0),(-1,-1),3),
        ("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]
    if header:
        style += [
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EAF0F7")),
            ("FONTNAME",(0,0),(-1,0),bold),
            ("TEXTCOLOR",(0,0),(-1,0),colors.HexColor("#0F2747")),
        ]
    t.setStyle(TableStyle(style))
    return t

def generate_pdf_bytes(root,ticker,mode="investment"):
    """PDF dạng flow: tự dồn trang, không ép mỗi mục một trang."""
    d,row,peer_row,mna_row=_report_context(root,ticker)
    font,bold=_register_pdf_fonts()
    st=_pdf_styles(font,bold)
    # Styles optimized for continuous flow.
    h1=ParagraphStyle("flow_h1",parent=st["h1"],fontSize=14,leading=18,spaceBefore=6,spaceAfter=4,keepWithNext=True)
    h2=ParagraphStyle("flow_h2",parent=st["h2"],fontSize=12,leading=15,spaceBefore=5,spaceAfter=3,keepWithNext=True)
    body=ParagraphStyle("flow_body",parent=st["body"],fontSize=11,leading=14.3,spaceBefore=6,spaceAfter=0,alignment=TA_JUSTIFY)
    small=ParagraphStyle("flow_small",parent=st["small"],fontSize=10,leading=13,spaceAfter=3,alignment=TA_JUSTIFY)

    qa=assess_report_quality(row,d["cfg"],d["hist"],d["prices"])
    norm_flags=normalization_flags(row,d["cfg"],d["prices"])
    strategic_case=strategic_reasonableness(row,d["cfg"],d.get("research"))
    strategic_case_text=reasonableness_conclusion(strategic_case)
    allmean=all_bank_means(d["summary"])
    cats,risks=catalysts_risks(row)

    bio=BytesIO()
    doc=SimpleDocTemplate(
        bio,pagesize=A4,leftMargin=MARGIN,rightMargin=MARGIN,
        topMargin=22*mm,bottomMargin=14*mm,
        title=f"Báo cáo phân tích, định giá & M&A ngân hàng - {ticker}",
        author="Nền tảng Phân tích, Định giá & M&A Ngân hàng Việt Nam"
    )
    story=[]

    def P(text,style=body):
        return Paragraph(escape(str(text)).replace("\n","<br/>"),style)

    def section(title,text=None):
        story.append(P(title,h1))
        if text:
            story.append(P(text,body))

    def chart(bio_img):
        story.append(_rl_flow_image(bio_img,178*mm))
        story.append(Spacer(1,3*mm))

    # 1
    section(f"1. TÓM TẮT ĐIỀU HÀNH - {ticker}",executive_summary(row,peer_row))
    story.append(P(f"TRẠNG THÁI: {qa['ReportStamp']} | Độ phủ: {_vi_num(qa['ReportCoverage']*100,0)}%",body))
    kpi=[
        ["Chỉ tiêu","Giá trị","Chỉ tiêu","Giá trị"],
        ["Giá thị trường",money(row.get("Price")),"Giá trị cơ bản",money(row.get("FairValue_Base"))],
        ["Chênh lệch cơ bản",pct(row.get("Upside_Base")),"ROE",pct(row.get("ROE_Used"))],
        ["P/B",mult(row.get("PB_Current")),"NPL",pct(row.get("NPL"))],
    ]
    story.append(_flow_table(kpi,font,bold,[36*mm,48*mm,36*mm,48*mm],8))
    story.append(Spacer(1,3*mm))
    chart(_chart_scores(row,d["summary"]))
    if n(row.get("StrategicPriceLow")) is not None:
        source_label="Thông tin thị trường do người dùng cung cấp" if str(row.get("StrategicSource",""))=="USER_MARKET_INTELLIGENCE" else str(row.get("StrategicSource","N/A"))
        story.append(P(f"Thông tin thị trường/M&A: {money(row.get('StrategicPriceLow'))} - {money(row.get('StrategicPriceHigh'))}. Nguồn: {source_label}. Khoảng này không phải giá trị cơ bản.",small))

    # 2
    section("2. HỒ SƠ NGÂN HÀNG & VỊ THẾ TƯƠNG ĐỐI",business_profile(row,peer_row))
    chart(_chart_history(d["hist"],ticker,["TotalAssets"],"Tổng tài sản",False))
    chart(_chart_history(d["hist"],ticker,["GrossLoans"],"Dư nợ khách hàng",False))
    chart(_chart_history(d["hist"],ticker,["CustomerDeposits"],"Tiền gửi khách hàng",False))

    # 3
    section("3. KHẢ NĂNG SINH LỜI & HIỆU QUẢ HOẠT ĐỘNG",profitability_text(row,peer_row))
    for metric,title in [
        ("ROE","Tỷ suất sinh lời trên vốn chủ sở hữu (ROE)"),
        ("ROA","Tỷ suất sinh lời trên tổng tài sản (ROA)"),
        ("NIM","Biên lãi ròng (NIM)"),
        ("CIR","Tỷ lệ chi phí/thu nhập (CIR)")
    ]:
        chart(_chart_history(d["hist"],ticker,[metric],title,True))

    # 4
    section("4. CHẤT LƯỢNG TÀI SẢN & CHI PHÍ TÍN DỤNG",asset_quality_text(row,peer_row))
    chart(_chart_history(d["hist"],ticker,["NPL"],"Xu hướng tỷ lệ nợ xấu (NPL)",True))
    story.append(P("Lưu ý: chi phí dự phòng có thể được Vnstock/BCTC thể hiện theo đơn vị tiền tệ; cần đọc cùng nguồn dữ liệu gốc khi so sánh với tỷ lệ NPL.",small))

    # 5
    section("5. NGUỒN VỐN, THANH KHOẢN & CHI PHÍ HUY ĐỘNG",funding_text(row,peer_row))
    chart(_chart_history(d["hist"],ticker,["CASA"],"Tỷ lệ tiền gửi không kỳ hạn (CASA)",True))
    chart(_chart_history(d["hist"],ticker,["LDR"],"Tỷ lệ cho vay trên tiền gửi (LDR)",True))

    # 6
    section("6. VỐN, KHẢ NĂNG CHỐNG CHỊU & NĂNG LỰC TĂNG TRƯỞNG",capital_text(row,peer_row))
    chart(_chart_history(d["hist"],ticker,["CAR"],"Xu hướng hệ số an toàn vốn (CAR)",True))

    # 7
    section("7. ĐỊNH GIÁ TƯƠNG ĐỐI & NHÓM SO SÁNH",valuation_text(row,peer_row))
    chart(_chart_peer(d["summary"],ticker))
    if peer_row:
        pdata=[
            ["Chỉ tiêu",ticker,"Trung bình nhóm so sánh","Bình quân 20 ngân hàng niêm yết"],
            ["ROE",pct(row.get("ROE_Used")),pct(peer_row.get("MeanROE")),pct(allmean.get("ROE_Used"))],
            ["NIM",pct(row.get("NIM")),pct(peer_row.get("MeanNIM")),pct(allmean.get("NIM"))],
            ["NPL",pct(row.get("NPL")),pct(peer_row.get("MeanNPL")),pct(allmean.get("NPL"))],
            ["CAR",pct(row.get("CAR")),pct(peer_row.get("MeanCAR")),pct(allmean.get("CAR"))],
            ["CASA",pct(row.get("CASA")),pct(peer_row.get("MeanCASA")),pct(allmean.get("CASA"))],
            ["P/B",mult(row.get("PB_Current")),mult(peer_row.get("MeanPB")),mult(allmean.get("PB_Current"))]
        ]
        story.append(_flow_table(pdata,font,bold,[30*mm,35*mm,52*mm,52*mm],7.4))
        story.append(Spacer(1,3*mm))

    # 8
    if mode=="mna":
        section("8. GIÁ TRỊ QUYỀN KIỂM SOÁT & MÔ PHỎNG M&A",mna_text(row,mna_row))
        if mna_row:
            data=[
                ["Chỉ tiêu","Giá trị"],
                ["Giá trị độc lập",bn(mna_row.get("StandaloneEquityValue"))],
                ["Giá trị hiện tại của cộng hưởng",bn(mna_row.get("PV_Synergies"))],
                ["Chi phí tích hợp",bn(mna_row.get("IntegrationCost"))],
                ["Giá trị thanh toán",bn(mna_row.get("IllustrativeConsideration"))],
                ["P/B giao dịch",mult(mna_row.get("ImpliedPB"))],
                ["P/TBV giao dịch",mult(mna_row.get("ImpliedPTBV"))],
            ]
            story.append(_flow_table(data,font,bold,[90*mm,80*mm],8))
    else:
        section("8. ĐỊNH GIÁ CƠ BẢN & GIÁ TRỊ CHIẾN LƯỢC",valuation_text(row,peer_row))
        if n(row.get("StrategicPriceLow")) is not None:
            story.append(P(
                f"Lớp giá trị chiến lược/M&A được ghi nhận riêng ở {money(row.get('StrategicPriceLow'))} - {money(row.get('StrategicPriceHigh'))}. "
                f"Nguồn: {row.get('StrategicSource','N/A')}; ngày {row.get('StrategicAsOfDate','N/A')}; mức độ tin cậy {row.get('StrategicConfidence','N/A')}. "
                "Khoảng này không được dùng để ép ngược giá trị cơ bản; chênh lệch cần được giải thích bằng quy mô lô, quyền kiểm soát, tính khan hiếm, tái cơ cấu và giá trị cộng hưởng.",
                body
            ))
        chart(_chart_valuation(d["methods"],row,d["summary"]))

    # 9
    if mode=="mna":
        section("9. PPA, LỢI THẾ THƯƠNG MẠI & TÁC ĐỘNG TÁI CƠ CẤU",
                "Trong phương pháp mua lại, giá mua được phân bổ vào tài sản và nợ phải trả có thể xác định theo giá trị hợp lý. "
                "Đối với ngân hàng, điều chỉnh rủi ro tín dụng danh mục cho vay, chứng khoán, nghĩa vụ ngoại bảng và tài sản vô hình liên quan đến khách hàng/tiền gửi cần được thẩm định riêng. "
                "Lợi thế thương mại chỉ là phần chênh lệch còn lại. Nếu ngân hàng mục tiêu cần bổ sung vốn, phần vốn bổ sung phải được tách khỏi giá mua tối đa.")
        chart(_chart_sensitivity(row))
    else:
        section("9. KIỂM TRA TÍNH HỢP LÝ 80.000-100.000 ĐỒNG/CP",
                "Đối chiếu định lượng vùng giá chiến lược với thị giá, BVPS hiện tại/hậu xử lý, giá trị lô 32,5%, ước tính 63.250 tỷ đồng gốc+lãi liên quan và các nghiên cứu công khai. Đây là kiểm tra tính hợp lý kinh tế, không phải khẳng định giá giao dịch tương lai.")
        if strategic_case.get("strategic_low") is not None:
            lo=strategic_case["low"]; hi=strategic_case["high"]
            data=[
                ["Chỉ tiêu","80.000 đồng/cp","100.000 đồng/cp"],
                ["Mức cao hơn thị giá",pct(lo.get("premium_to_market")),pct(hi.get("premium_to_market"))],
                ["P/B trên BVPS hiện tại",mult(lo.get("implied_pb_current")),mult(hi.get("implied_pb_current"))],
                ["P/B hậu xử lý (BVPS +10.000)",mult(lo.get("implied_pb_post_resolution")),mult(hi.get("implied_pb_post_resolution"))],
                ["Giá trị lô 32,5%",bn((lo.get("block_consideration_bn") or 0)*1e9),bn((hi.get("block_consideration_bn") or 0)*1e9)],
                ["Bao phủ 63.250 tỷ gốc+lãi",pct(lo.get("claim_recovery")),pct(hi.get("claim_recovery"))],
                ["So với kịch bản cao công khai",pct(lo.get("premium_to_public_high")),pct(hi.get("premium_to_public_high"))],
            ]
            story.append(_flow_table(data,font,bold,[75*mm,48*mm,48*mm],7.4))
            story.append(Spacer(1,3*mm))
            story.append(P(
                "80.000 đồng/cp là mức chiến lược tương đối dễ biện minh vì chỉ cao hơn nhẹ so với thị giá và nằm sát vùng kịch bản cao của định giá độc lập công khai. "
                "100.000 đồng/cp là mức hợp lý có điều kiện: giá trị lô 32,5% tiến gần mức thu hồi đầy đủ 63.250 tỷ đồng, vì vậy cần xác suất xử lý cao cùng quyền ảnh hưởng và giá trị hậu tái cơ cấu.",
                small
            ))
        chart(_chart_sensitivity(row))

    # 10
    section("10. KẾT LUẬN, ĐỘNG LỰC & RỦI RO",executive_summary(row,peer_row))
    story.append(P("ĐỘNG LỰC TIỀM NĂNG",h2))
    for x in cats[:5]: story.append(P("• "+x,body))
    story.append(P("RỦI RO CHÍNH",h2))
    for x in risks[:5]: story.append(P("• "+x,body))
    story.append(P("ĐIỂM CẦN LƯU Ý KHI ĐỊNH GIÁ",h2))
    for x in norm_flags[:5]: story.append(P("• "+x,small))
    story.append(P(
        "Ghi chú phương pháp: Thu nhập thặng dư là phương pháp nội tại chính; P/B hợp lý, P/B nhóm so sánh và P/B lịch sử là các phép kiểm tra chéo. "
        "Trong M&A, giá trị quyền kiểm soát, giá trị cộng hưởng, PPA và nhu cầu bổ sung vốn được tách khỏi giá trị cơ bản độc lập. "
        "Mọi giả định phải được gắn nhãn và không được trình bày như dữ liệu quan sát thực tế.",small
    ))

    decor=lambda canv,docobj:_pdf_page_decor(canv,docobj,ticker,font,bold,qa["ReportStamp"])
    doc.build(story,onFirstPage=decor,onLaterPages=decor)
    bio.seek(0)
    return bio.getvalue()

def _doc_font(doc):
    styles=doc.styles
    for style_name in ["Normal","Title","Heading 1","Heading 2"]:
        style=styles[style_name]
        style.font.name="Lato"
        style._element.rPr.rFonts.set(qn("w:ascii"),"Lato")
        style._element.rPr.rFonts.set(qn("w:hAnsi"),"Lato")
        style._element.rPr.rFonts.set(qn("w:eastAsia"),"Lato")
        style._element.rPr.rFonts.set(qn("w:cs"),"Lato")

    normal=styles["Normal"]
    normal.font.size=Pt(11)
    normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before=Pt(6)
    normal.paragraph_format.space_after=Pt(0)
    normal.paragraph_format.line_spacing=1.3

    styles["Heading 1"].font.size=Pt(14)
    styles["Heading 1"].font.bold=True
    styles["Heading 1"].font.color.rgb=None
    styles["Heading 1"].paragraph_format.space_before=Pt(8)
    styles["Heading 1"].paragraph_format.space_after=Pt(4)
    styles["Heading 1"].paragraph_format.keep_with_next=True

    styles["Heading 2"].font.size=Pt(12)
    styles["Heading 2"].font.bold=True
    styles["Heading 2"].paragraph_format.space_before=Pt(6)
    styles["Heading 2"].paragraph_format.space_after=Pt(3)
    styles["Heading 2"].paragraph_format.keep_with_next=True

    # Bullet/list text follows the 11pt Lato body standard.
    for style_name in ["List Bullet","List Number"]:
        if style_name in styles:
            s=styles[style_name]
            s.font.name="Lato"
            s._element.rPr.rFonts.set(qn("w:ascii"),"Lato")
            s._element.rPr.rFonts.set(qn("w:hAnsi"),"Lato")
            s._element.rPr.rFonts.set(qn("w:eastAsia"),"Lato")
            s.font.size=Pt(11)
            s.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            s.paragraph_format.line_spacing=1.3


def _add_doc_header(section,ticker,page):
    header=section.header.paragraphs[0]
    header.text=f"BÁO CÁO PHÂN TÍCH, ĐỊNH GIÁ & M&A NGÂN HÀNG | {ticker} | ENGINE V6.4"
    for run in header.runs:
        run.font.name="Lato"
        run.font.size=Pt(10)
    footer=section.footer.paragraphs[0]
    footer.text="Tài liệu phân tích - không phải khuyến nghị đầu tư/chào mua."
    for run in footer.runs:
        run.font.name="Lato"
        run.font.size=Pt(10)

def _add_picture(doc,bio,width=178):
    bio.seek(0)
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2)
    p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.keep_together=True
    run=p.add_run()
    run.add_picture(bio,width=Mm(width))
    return p

def _normalize_docx_typography(doc):
    """Enforce final report typography after all content has been created."""
    # Main narrative paragraphs: Lato 11, justified. Keep headings and image paragraphs centered/as styled.
    for p in doc.paragraphs:
        style_name = p.style.name if p.style is not None else ""
        has_drawing = bool(p._p.xpath(".//w:drawing"))
        if style_name not in ("Heading 1","Heading 2","Title") and not has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name="Lato"
            r._element.rPr.rFonts.set(qn("w:ascii"),"Lato")
            r._element.rPr.rFonts.set(qn("w:hAnsi"),"Lato")
            r._element.rPr.rFonts.set(qn("w:eastAsia"),"Lato")
            r._element.rPr.rFonts.set(qn("w:cs"),"Lato")
            if style_name not in ("Heading 1","Heading 2","Title"):
                r.font.size=Pt(11)

    # Tables: Lato 10 throughout; justified cell text.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing=1.15
                    p.paragraph_format.space_before=Pt(0)
                    p.paragraph_format.space_after=Pt(0)
                    for r in p.runs:
                        r.font.name="Lato"
                        r._element.rPr.rFonts.set(qn("w:ascii"),"Lato")
                        r._element.rPr.rFonts.set(qn("w:hAnsi"),"Lato")
                        r._element.rPr.rFonts.set(qn("w:eastAsia"),"Lato")
                        r._element.rPr.rFonts.set(qn("w:cs"),"Lato")
                        r.font.size=Pt(10)

    # Header/footer: Lato 10.
    for sec in doc.sections:
        for p in list(sec.header.paragraphs)+list(sec.footer.paragraphs):
            for r in p.runs:
                r.font.name="Lato"
                r._element.rPr.rFonts.set(qn("w:ascii"),"Lato")
                r._element.rPr.rFonts.set(qn("w:hAnsi"),"Lato")
                r._element.rPr.rFonts.set(qn("w:eastAsia"),"Lato")
                r.font.size=Pt(10)

def generate_docx_bytes(root,ticker,mode="investment"):
    d,row,peer_row,mna_row=_report_context(root,ticker)
    qa=assess_report_quality(row,d["cfg"],d["hist"],d["prices"])
    norm_flags=normalization_flags(row,d["cfg"],d["hist"] if False else d["prices"])
    strategic_case=strategic_reasonableness(row,d["cfg"],d.get("research"))
    strategic_case_text=reasonableness_conclusion(strategic_case)
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Mm(210); sec.page_height=Mm(297); sec.top_margin=Mm(16); sec.bottom_margin=Mm(15); sec.left_margin=Mm(16); sec.right_margin=Mm(16)
    _doc_font(doc); cats,risks=catalysts_risks(row)

    pages=[
        ("1. TÓM TẮT ĐIỀU HÀNH",executive_summary(row,peer_row),_chart_scores(row,d["summary"])),
        ("2. HỒ SƠ NGÂN HÀNG & VỊ THẾ TƯƠNG ĐỐI",business_profile(row,peer_row),[
            _chart_history(d["hist"],ticker,["TotalAssets"],"Tổng tài sản",False),
            _chart_history(d["hist"],ticker,["GrossLoans"],"Dư nợ khách hàng",False),
            _chart_history(d["hist"],ticker,["CustomerDeposits"],"Tiền gửi khách hàng",False)]),
        ("3. KHẢ NĂNG SINH LỜI & HIỆU QUẢ HOẠT ĐỘNG",profitability_text(row,peer_row),[
            _chart_history(d["hist"],ticker,["ROE"],"Tỷ suất sinh lời trên vốn chủ sở hữu (ROE)",True),
            _chart_history(d["hist"],ticker,["ROA"],"Tỷ suất sinh lời trên tổng tài sản (ROA)",True),
            _chart_history(d["hist"],ticker,["NIM"],"Biên lãi ròng (NIM)",True),
            _chart_history(d["hist"],ticker,["CIR"],"Tỷ lệ chi phí/thu nhập (CIR)",True)]),
        ("4. CHẤT LƯỢNG TÀI SẢN & CHI PHÍ TÍN DỤNG",asset_quality_text(row,peer_row),_chart_history(d["hist"],ticker,["NPL"],"Xu hướng tỷ lệ nợ xấu (NPL)",True)),
        ("5. NGUỒN VỐN, THANH KHOẢN & CHI PHÍ HUY ĐỘNG",funding_text(row,peer_row),[
            _chart_history(d["hist"],ticker,["CASA"],"Tỷ lệ tiền gửi không kỳ hạn (CASA)",True),
            _chart_history(d["hist"],ticker,["LDR"],"Tỷ lệ cho vay trên tiền gửi (LDR)",True)]),
        ("6. VỐN, KHẢ NĂNG CHỐNG CHỊU & NĂNG LỰC TĂNG TRƯỞNG",capital_text(row,peer_row),_chart_history(d["hist"],ticker,["CAR"],"Xu hướng hệ số an toàn vốn (CAR)",True)),
        ("7. ĐỊNH GIÁ TƯƠNG ĐỐI & NHÓM SO SÁNH",valuation_text(row,peer_row),_chart_peer(d["summary"],ticker)),
        ("8. GIÁ TRỊ QUYỀN KIỂM SOÁT & MÔ PHỎNG M&A" if mode=="mna" else "8. ĐỊNH GIÁ CƠ BẢN & KIỂM TRA GIÁ THÂU TÓM",mna_text(row,mna_row) if mode=="mna" else valuation_text(row,peer_row)+" "+strategic_case_text,_chart_valuation(d["methods"],row,d["summary"])),
        ("9. PPA, GOODWILL & TÁC ĐỘNG TÁI CƠ CẤU" if mode=="mna" else "9. KIỂM TRA TÍNH HỢP LÝ 80.000–100.000 ĐỒNG/CP","Đối chiếu định lượng vùng giá chiến lược với thị giá, BVPS hiện tại/hậu xử lý, giá trị lô 32,5%, ước tính 63.250 tỷ đồng gốc+lãi liên quan và các benchmark nghiên cứu công khai. Đây là kiểm tra tính hợp lý kinh tế, không phải khẳng định giá giao dịch tương lai.",_chart_sensitivity(row)),
        ("10. KẾT LUẬN, ĐỘNG LỰC & RỦI RO",executive_summary(row,peer_row),None),
    ]
    for i,(head,body,chart) in enumerate(pages,1):
        _add_doc_header(doc.sections[-1],ticker,i)
        p=doc.add_paragraph(); p.style="Heading 1"; p.paragraph_format.keep_with_next=True; p.add_run(head)
        body_p=doc.add_paragraph(body)
        body_p.paragraph_format.widow_control=True
        if i==1:
            p=doc.add_paragraph()
            r=p.add_run(f"TRẠNG THÁI: {qa['ReportStamp']} | Độ phủ: {qa['ReportCoverage']:.0%}")
            r.bold=True
            table=doc.add_table(rows=3,cols=4); table.style="Table Grid"
            vals=[("Giá thị trường",money(row.get("Price"))),("Giá trị cơ bản",money(row.get("FairValue_Base"))),("Chênh lệch cơ bản",pct(row.get("Upside_Base"))),("ROE",pct(row.get("ROE_Used"))),("P/B",mult(row.get("PB_Current"))),("NPL",pct(row.get("NPL")))]
            for j,(k,v) in enumerate(vals):
                rr=j//2; cc=(j%2)*2; table.cell(rr,cc).text=k; table.cell(rr,cc+1).text=v
        if chart is not None:
            if isinstance(chart,(list,tuple)):
                # V6.3: mọi biểu đồ dùng gần toàn bộ chiều ngang vùng in A4.
                # Word tự đẩy biểu đồ xuống trang kế tiếp nếu phần còn lại không đủ chỗ.
                for ch in chart:
                    _add_picture(doc,ch,178)
            else:
                _add_picture(doc,chart,178)
        if i==10:
            doc.add_paragraph("ĐỘNG LỰC TIỀM NĂNG",style="Heading 2")
            for x in cats[:5]: doc.add_paragraph(x,style="List Bullet")
            doc.add_paragraph("RỦI RO CHÍNH",style="Heading 2")
            for x in risks[:5]: doc.add_paragraph(x,style="List Bullet")
            doc.add_paragraph("ĐIỂM CẦN LƯU Ý KHI ĐỊNH GIÁ",style="Heading 2")
            for x in norm_flags[:5]: doc.add_paragraph(x,style="List Bullet")
            # Disclaimer đã có ở footer; không lặp lại để tránh tạo một trang trắng gần như hoàn toàn.
        # Không chèn page break cưỡng bức: Word tự dồn trang theo nội dung thực tế.
    _normalize_docx_typography(doc)
    bio=BytesIO(); doc.save(bio); return bio.getvalue()
