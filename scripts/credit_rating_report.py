from __future__ import annotations
from pathlib import Path
from io import BytesIO
import math, json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Image as RLImage
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from xml.sax.saxutils import escape

try:
    from scripts.credit_rating_engine import build_credit_rating, FACTOR_LABELS
except Exception:
    from credit_rating_engine import build_credit_rating, FACTOR_LABELS

PAGE_W,PAGE_H=A4; MARGIN=16*mm
plt.rcParams.update({"font.family":"Lato","font.size":10,"axes.titlesize":10,"axes.labelsize":10,"xtick.labelsize":10,"ytick.labelsize":10,"legend.fontsize":10})

def _csv(p):
    try:return pd.read_csv(p)
    except Exception:return pd.DataFrame()

def _load(root):
    root=Path(root); data=root/'data'; out=data/'model_outputs'
    return _csv(out/'valuation_summary.csv'),_csv(data/'bank_history_long.csv')

def _n(x):
    try:
        v=float(x); return v if math.isfinite(v) else None
    except:return None

def _vi(x,d=1):
    v=_n(x)
    if v is None:return 'N/A'
    s=f"{v:,.{d}f}"; return s.replace(',','§').replace('.',',').replace('§','.')

def _pct(x,d=1):return 'N/A' if _n(x) is None else _vi(_n(x)*100,d)+'%'
def _bn(x):return 'N/A' if _n(x) is None else _vi(_n(x)/1e9,0)+' tỷ đồng'

def _font_paths():
    candidates=[('/usr/share/fonts/truetype/lato/Lato-Regular.ttf','/usr/share/fonts/truetype/lato/Lato-Bold.ttf'),('C:/Windows/Fonts/Lato-Regular.ttf','C:/Windows/Fonts/Lato-Bold.ttf')]
    try:
        from matplotlib import font_manager
        candidates.insert(0,(font_manager.findfont('Lato',fallback_to_default=False),font_manager.findfont(font_manager.FontProperties(family='Lato',weight='bold'),fallback_to_default=False)))
    except:pass
    for a,b in candidates:
        if Path(a).exists() and Path(b).exists():return a,b
    return None,None

def _reg_fonts():
    a,b=_font_paths()
    if not a:return 'Helvetica','Helvetica-Bold'
    if 'CRLato' not in pdfmetrics.getRegisteredFontNames():pdfmetrics.registerFont(TTFont('CRLato',a))
    if 'CRLato-Bold' not in pdfmetrics.getRegisteredFontNames():pdfmetrics.registerFont(TTFont('CRLato-Bold',b))
    return 'CRLato','CRLato-Bold'

def _chart_factor(rating):
    fs=rating['FactorScores']; labels=[FACTOR_LABELS[k] for k in fs]; vals=[]
    for k in fs:
        v=fs[k]
        vals.append(v if k not in {'Funding','Liquidity'} else 1+(v-1)*5/3)
    fig,ax=plt.subplots(figsize=(8.2,3.0)); ax.barh(labels,vals); ax.invert_yaxis(); ax.set_xlim(.7,6.3); ax.set_xlabel('Điểm yếu tố (1 = mạnh nhất; 6 = yếu nhất)'); ax.set_title('Khung yếu tố xếp hạng tín nhiệm'); ax.grid(axis='x',alpha=.2)
    bio=BytesIO(); fig.tight_layout(); fig.savefig(bio,format='png',dpi=160,bbox_inches='tight'); plt.close(fig); bio.seek(0); return bio

def _period_date(s):
    import re
    m=re.search(r'(20\d{2}).*?Q\s*([1-4])',str(s).upper())
    if not m:return pd.NaT
    return pd.Timestamp(int(m.group(1)),(int(m.group(2))-1)*3+1,1)

def _chart_metric(hist,summary,ticker,metric,title):
    h=hist[hist.Ticker.astype(str).eq(str(ticker)) & hist.Metric.astype(str).eq(metric)].copy()
    p=hist[hist.Metric.astype(str).eq(metric)].copy()
    for x in (h,p):
        x['Value']=pd.to_numeric(x.Value,errors='coerce'); x['Date']=x.Period.map(_period_date)
    h=h.dropna(subset=['Value','Date']).sort_values('Date'); p=p.dropna(subset=['Value','Date'])
    if metric in {'ROE','ROA','NIM','NPL','CAR','CIR','LDR','CASA'}:
        h=h[h.Value!=0]; p=p[p.Value!=0]
    if metric=='CAR':h=h[h.Value>0];p=p[p.Value>0]
    pm=p.groupby('Date',as_index=False).Value.mean().sort_values('Date')
    fig,ax=plt.subplots(figsize=(8.2,2.85));
    if len(h):ax.plot(h.Date,h.Value,marker='o',label=ticker)
    if len(pm):ax.plot(pm.Date,pm.Value,linestyle='--',label='Bình quân 20 ngân hàng niêm yết')
    ax.set_title(title+' - so với bình quân 20 ngân hàng niêm yết');ax.grid(alpha=.2);ax.legend();
    if metric in {'ROE','ROA','NIM','NPL','CAR','CIR','LDR','CASA'}:ax.yaxis.set_major_formatter(FuncFormatter(lambda v,pos:_vi(v*100,1)+'%'))
    fig.autofmt_xdate();bio=BytesIO();fig.tight_layout();fig.savefig(bio,format='png',dpi=160,bbox_inches='tight');plt.close(fig);bio.seek(0);return bio

def _rl_img(bio,w=178*mm):
    from PIL import Image
    bio.seek(0); im=Image.open(BytesIO(bio.getvalue())); ratio=im.height/im.width; bio.seek(0);return RLImage(bio,width=w,height=w*ratio)

def _table(data,font,bold,widths=None,fs=10):
    t=Table(data,colWidths=widths,repeatRows=1,hAlign='LEFT')
    t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.35,colors.HexColor('#CBD5E1')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#EAF0F7')),('FONTNAME',(0,0),(-1,-1),font),('FONTNAME',(0,0),(-1,0),bold),('FONTSIZE',(0,0),(-1,-1),fs),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4),('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4)]));return t

def generate_credit_rating_pdf(root,ticker,governance_score=3,external_support_notches=0,analyst_notches=0,notch_overrides=None):
    summary,hist=_load(root); row=summary[summary.Ticker.astype(str).eq(str(ticker))].iloc[0]; cr=build_credit_rating(summary,ticker,governance_score,external_support_notches,analyst_notches,notch_overrides)
    font,bold=_reg_fonts(); bio=BytesIO(); doc=SimpleDocTemplate(bio,pagesize=A4,leftMargin=MARGIN,rightMargin=MARGIN,topMargin=18*mm,bottomMargin=15*mm,title=f'Báo cáo xếp hạng tín nhiệm mô phỏng - {ticker}')
    h1=ParagraphStyle('h1',fontName=bold,fontSize=14,leading=18,textColor=colors.HexColor('#0F2747'),spaceBefore=6,spaceAfter=4)
    h2=ParagraphStyle('h2',fontName=bold,fontSize=12,leading=15,textColor=colors.HexColor('#2F6B2F'),spaceBefore=5,spaceAfter=3)
    body=ParagraphStyle('body',fontName=font,fontSize=11,leading=14.3,alignment=TA_JUSTIFY,spaceBefore=6,spaceAfter=0)
    small=ParagraphStyle('small',fontName=font,fontSize=10,leading=13,alignment=TA_JUSTIFY,textColor=colors.HexColor('#555555'))
    story=[]
    P=lambda x,s=body:Paragraph(escape(str(x)).replace('\n','<br/>'),s)
    story += [P(f'BÁO CÁO XẾP HẠNG TÍN NHIỆM NGÂN HÀNG - {ticker}',h1),P('KẾT QUẢ MÔ PHỎNG / NỘI BỘ',h2)]
    story.append(_table([['BICRA / Anchor','SACP','ICR cuối cùng','Triển vọng'],[cr['AnchorRating'],cr['SACPRating'],cr['FinalRating'],cr['Outlook']]],font,bold,[44*mm,44*mm,44*mm,42*mm],10))
    story += [Spacer(1,3*mm),P('LUẬN ĐIỂM XẾP HẠNG',h2),P(f"Theo methodology 2025, điểm xuất phát của các ngân hàng Việt Nam là BICRA/Anchor {cr['AnchorRating']}. Sau khi cộng/trừ {cr['InternalNotches']:+d} notch từ Hồ sơ kinh doanh, Vốn & lợi nhuận, Vị thế rủi ro, ma trận Huy động vốn - Thanh khoản và các điều chỉnh khác, {ticker} có SACP {cr['SACPRating']}. Điều chỉnh hỗ trợ bên ngoài {cr['ExternalSupportNotches']:+d} notch đưa ICR mô phỏng về {cr['FinalRating']}, triển vọng {cr['Outlook']}.")]
    story += [_rl_img(_chart_factor(cr)),P('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG',h2),P('Điểm mạnh',h2)]
    for x in cr['Strengths']:story.append(P('• '+x))
    story.append(P('Điểm hạn chế',h2))
    for x in cr['Constraints']:story.append(P('• '+x))
    story += [P('TRIỂN VỌNG',h2),P(f"Triển vọng {cr['Outlook']} phản ánh cân bằng giữa khả năng sinh lời, đệm vốn, chất lượng tài sản và cấu trúc huy động hiện tại. Mô hình không tự coi tăng trưởng cao là tích cực nếu tăng trưởng tín dụng vượt đáng kể huy động hoặc đi kèm suy giảm chất lượng tài sản.")]
    story.append(P('CÁC YẾU TỐ CÓ THỂ DẪN ĐẾN NÂNG BẬC',h2)); [story.append(P('• '+x)) for x in cr['UpgradeTriggers']]
    story.append(P('CÁC YẾU TỐ CÓ THỂ DẪN ĐẾN HẠ BẬC',h2)); [story.append(P('• '+x)) for x in cr['DowngradeTriggers']]
    story += [P('1. THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH',h1),P(f"{ticker} được đặt trong bộ so sánh 20 ngân hàng niêm yết. Tổng tài sản gần nhất khoảng {_bn(row.get('TotalAssets'))}; vốn chủ sở hữu {_bn(row.get('Equity'))}; dư nợ khách hàng {_bn(row.get('GrossLoans'))}; tiền gửi khách hàng {_bn(row.get('CustomerDeposits'))}.")]
    story += [P('2. RỦI RO VĨ MÔ & NGÀNH NGÂN HÀNG',h1),P('Môi trường vĩ mô Việt Nam tiếp tục tạo nền tảng tương đối thuận lợi cho ngành ngân hàng nhờ tăng trưởng kinh tế cao, đầu tư và tiêu dùng cải thiện, lạm phát được kiểm soát và chính sách tiền tệ duy trì linh hoạt theo hướng hỗ trợ tăng trưởng. Năm 2025, tăng trưởng tín dụng toàn hệ thống khoảng 19%, trong khi CPI bình quân tăng 3,31%; thị trường tiền tệ và ngoại hối nhìn chung ổn định. Tuy nhiên, bất định thương mại toàn cầu, áp lực tỷ giá và lạm phát làm dư địa nới lỏng tiền tệ bị giới hạn. Đối với ngành ngân hàng, tín dụng tăng 19,01% cao hơn đáng kể tăng trưởng huy động 14,11%, tạo áp lực lên thanh khoản và chi phí vốn; khả năng mở rộng NIM bị hạn chế do lãi suất cho vay vẫn cần hỗ trợ nền kinh tế. Tỷ lệ tín dụng/GDP của Việt Nam ở mức cao, trong khi hệ số an toàn vốn trung bình của hệ thống còn tương đối thấp so với một số nước trong khu vực. Bước sang 2026, chất lượng tăng trưởng tín dụng, khả năng hấp thụ vốn và dòng tiền của khách hàng, cơ cấu nguồn vốn, chất lượng tài sản và năng lực quản trị rủi ro là các biến số chính. Theo methodology đã phê duyệt, tổng hợp rủi ro kinh tế vĩ mô và rủi ro ngành cho BICRA Việt Nam ở mức a-, được sử dụng làm Anchor ban đầu cho ngân hàng.')]
    sections=[('3. VỊ THẾ KINH DOANH','BusinessPosition',[('TotalAssets','Tổng tài sản'),('GrossLoans','Dư nợ khách hàng'),('CustomerDeposits','Tiền gửi khách hàng')]),('4. VỐN, ĐÒN BẨY & LỢI NHUẬN','CapitalEarnings',[('CAR','Hệ số an toàn vốn (CAR)'),('ROE','ROE'),('NIM','Biên lãi ròng (NIM)')]),('5. VỊ THẾ RỦI RO','RiskPosition',[('NPL','Tỷ lệ nợ xấu (NPL)')]),('6. HUY ĐỘNG VỐN','Funding',[('CASA','Tỷ lệ tiền gửi không kỳ hạn (CASA)')]),('7. THANH KHOẢN','Liquidity',[('LDR','Tỷ lệ cho vay trên tiền gửi (LDR)')])]
    for title,key,charts in sections:
        story += [P(title,h1),P(f"Đánh giá: {cr['FactorScores'][key]}/"+("4" if key in {'Funding','Liquidity'} else "6")+f". {cr['FactorRationale'][key]}")]
        for metric,ct in charts:story += [_rl_img(_chart_metric(hist,summary,ticker,metric,ct)),Spacer(1,2*mm)]
    story += [P('8. CẦU NỐI BICRA - SACP - ICR',h1)]
    bridge=[['Bước','Điểm/Notch','Kết quả'],
            ['BICRA / Anchor','Điểm xuất phát',cr['AnchorRating']],
            ['Hồ sơ kinh doanh',f"{cr['BusinessNotch']:+d} notch",''],
            ['Vốn & lợi nhuận',f"{cr['CapitalNotch']:+d} notch",''],
            ['Vị thế rủi ro',f"{cr['RiskNotch']:+d} notch",''],
            ['Huy động vốn × Thanh khoản',f"{cr['FundingLiquidityNotch']:+d} notch",''],
            ['Điều chỉnh khác trước SACP',f"{cr['OtherInternalNotches']:+d} notch",''],
            ['SACP','',cr['SACPRating']],
            ['Hỗ trợ bên ngoài',f"{cr['ExternalSupportNotches']:+d} notch",''],
            ['ICR cuối cùng','',cr['FinalRating']]]
    story += [_table(bridge,font,bold,[70*mm,48*mm,52*mm],10),Spacer(1,3*mm),
              P('9. HỖ TRỢ BÊN NGOÀI',h1),P(f"Hỗ trợ bên ngoài được xem xét sau khi xác định SACP. Điều chỉnh hiện tại là {cr['ExternalSupportNotches']:+d} notch. Theo methodology, hỗ trợ Chính phủ/NHNN tập trung rõ hơn đối với nhóm ngân hàng quốc doanh; đối với ngân hàng TMCP tư nhân cần đánh giá theo từng trường hợp về khả năng và mức độ sẵn sàng hỗ trợ."),
              P('10. BẢNG ĐIỂM & PHƯƠNG PHÁP',h1)]
    factor_rows=[['Yếu tố','Điểm','Notch','Diễn giải']]
    for k,v in cr['FactorScores'].items():
        notch=cr['FactorNotches'].get(k,'—') if k not in {'Funding','Liquidity'} else '—'
        factor_rows.append([FACTOR_LABELS[k],f"{v}/"+('4' if k in {'Funding','Liquidity'} else '6'),str(notch),cr['FactorRationale'][k]])
    factor_rows.append(['Huy động vốn × Thanh khoản','Ma trận Bảng 10',str(cr['FundingLiquidityNotch']),f"{cr['FundingDescriptor']} × {cr['LiquidityDescriptor']}"])
    story += [_table(factor_rows,font,bold,[35*mm,22*mm,18*mm,95*mm],8.5),Spacer(1,3*mm),P('GIẢ ĐỊNH & HẠN CHẾ',h2),P(cr['Disclaimer'],small)]
    def decor(c,d):
        c.saveState();c.setFont(bold,8);c.setFillColor(colors.HexColor('#0F2747'));c.drawString(MARGIN,PAGE_H-10*mm,'BÁO CÁO XẾP HẠNG TÍN NHIỆM NGÂN HÀNG - MÔ PHỎNG/NỘI BỘ');c.drawRightString(PAGE_W-MARGIN,PAGE_H-10*mm,f'{ticker} | Trang {d.page}');c.setFont(font,7);c.setFillColor(colors.HexColor('#666666'));c.drawString(MARGIN,7*mm,'Nguồn dữ liệu: Vnstock/BCTC và mô hình nội bộ; cần chuyên viên rà soát trước khi sử dụng.');c.restoreState()
    doc.build(story,onFirstPage=decor,onLaterPages=decor);bio.seek(0);return bio.getvalue()

def _set_lato(run,size=None,bold=None):
    run.font.name='Lato'; run._element.rPr.rFonts.set(qn('w:ascii'),'Lato');run._element.rPr.rFonts.set(qn('w:hAnsi'),'Lato');run._element.rPr.rFonts.set(qn('w:eastAsia'),'Lato')
    if size:run.font.size=Pt(size)
    if bold is not None:run.bold=bold

def generate_credit_rating_docx(root,ticker,governance_score=3,external_support_notches=0,analyst_notches=0,notch_overrides=None):
    summary,hist=_load(root); row=summary[summary.Ticker.astype(str).eq(str(ticker))].iloc[0];cr=build_credit_rating(summary,ticker,governance_score,external_support_notches,analyst_notches,notch_overrides)
    doc=Document();sec=doc.sections[0];sec.page_width=Mm(210);sec.page_height=Mm(297);sec.top_margin=Mm(16);sec.bottom_margin=Mm(15);sec.left_margin=Mm(16);sec.right_margin=Mm(16)
    for sty in ['Normal','Title','Heading 1','Heading 2']:
        s=doc.styles[sty];s.font.name='Lato';s._element.rPr.rFonts.set(qn('w:ascii'),'Lato');s._element.rPr.rFonts.set(qn('w:hAnsi'),'Lato');s._element.rPr.rFonts.set(qn('w:eastAsia'),'Lato')
    doc.styles['Normal'].font.size=Pt(11);doc.styles['Normal'].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;doc.styles['Normal'].paragraph_format.line_spacing=1.3;doc.styles['Normal'].paragraph_format.space_before=Pt(6);doc.styles['Normal'].paragraph_format.space_after=Pt(0)
    doc.styles['Heading 1'].font.size=Pt(14);doc.styles['Heading 2'].font.size=Pt(12)
    hp=sec.header.paragraphs[0];hp.text=f'BÁO CÁO XẾP HẠNG TÍN NHIỆM NGÂN HÀNG | {ticker} | MÔ PHỎNG/NỘI BỘ';[_set_lato(r,10,True) for r in hp.runs]
    fp=sec.footer.paragraphs[0];fp.text='Nguồn: Vnstock/BCTC + mô hình nội bộ. Không phải kết quả xếp hạng tín nhiệm chính thức.';[_set_lato(r,9) for r in fp.runs]
    doc.add_heading(f'BÁO CÁO XẾP HẠNG TÍN NHIỆM NGÂN HÀNG - {ticker}',0);doc.add_heading('KẾT QUẢ MÔ PHỎNG / NỘI BỘ',1)
    table=doc.add_table(rows=2,cols=4);table.style='Table Grid';hdr=['BICRA / Anchor','SACP','ICR cuối cùng','Triển vọng'];vals=[cr['AnchorRating'],cr['SACPRating'],cr['FinalRating'],cr['Outlook']]
    for j,x in enumerate(hdr):table.cell(0,j).text=x
    for j,x in enumerate(vals):table.cell(1,j).text=str(x)
    def addtxt(h,text):doc.add_heading(h,1);doc.add_paragraph(text)
    addtxt('LUẬN ĐIỂM XẾP HẠNG',f"Điểm xuất phát BICRA/Anchor của ngân hàng Việt Nam là {cr['AnchorRating']}. Sau {cr['InternalNotches']:+d} notch điều chỉnh nội tại, {ticker} có SACP {cr['SACPRating']}; sau {cr['ExternalSupportNotches']:+d} notch hỗ trợ bên ngoài, ICR mô phỏng là {cr['FinalRating']}, triển vọng {cr['Outlook']}.")
    doc.add_picture(_chart_factor(cr),width=Mm(178))
    doc.add_heading('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG',1);doc.add_heading('Điểm mạnh',2);[doc.add_paragraph(x,style='List Bullet') for x in cr['Strengths']];doc.add_heading('Điểm hạn chế',2);[doc.add_paragraph(x,style='List Bullet') for x in cr['Constraints']]
    doc.add_heading('TRIỂN VỌNG',1);doc.add_paragraph(f"Triển vọng {cr['Outlook']} phản ánh cân bằng giữa khả năng sinh lời, đệm vốn, chất lượng tài sản và cấu trúc huy động hiện tại.")
    doc.add_heading('Các yếu tố có thể dẫn đến Nâng Bậc Xếp hạng',2);[doc.add_paragraph(x,style='List Bullet') for x in cr['UpgradeTriggers']]
    doc.add_heading('Các yếu tố có thể dẫn đến Hạ Bậc Xếp hạng',2);[doc.add_paragraph(x,style='List Bullet') for x in cr['DowngradeTriggers']]
    addtxt('1. THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH',f"Tổng tài sản gần nhất khoảng {_bn(row.get('TotalAssets'))}; vốn chủ sở hữu {_bn(row.get('Equity'))}; dư nợ khách hàng {_bn(row.get('GrossLoans'))}; tiền gửi khách hàng {_bn(row.get('CustomerDeposits'))}.")
    addtxt('2. RỦI RO VĨ MÔ & NGÀNH NGÂN HÀNG','Môi trường vĩ mô Việt Nam tiếp tục tạo nền tảng tương đối thuận lợi cho ngành ngân hàng nhờ tăng trưởng kinh tế cao, đầu tư và tiêu dùng cải thiện, lạm phát được kiểm soát và chính sách tiền tệ duy trì linh hoạt theo hướng hỗ trợ tăng trưởng. Năm 2025, tăng trưởng tín dụng toàn hệ thống khoảng 19%, trong khi CPI bình quân tăng 3,31%; thị trường tiền tệ và ngoại hối nhìn chung ổn định. Tuy nhiên, bất định thương mại toàn cầu, áp lực tỷ giá và lạm phát làm dư địa nới lỏng tiền tệ bị giới hạn. Đối với ngành ngân hàng, tín dụng tăng 19,01% cao hơn đáng kể tăng trưởng huy động 14,11%, tạo áp lực lên thanh khoản và chi phí vốn; khả năng mở rộng NIM bị hạn chế do lãi suất cho vay vẫn cần hỗ trợ nền kinh tế. Tỷ lệ tín dụng/GDP của Việt Nam ở mức cao, trong khi hệ số an toàn vốn trung bình của hệ thống còn tương đối thấp so với một số nước trong khu vực. Bước sang 2026, chất lượng tăng trưởng tín dụng, khả năng hấp thụ vốn và dòng tiền của khách hàng, cơ cấu nguồn vốn, chất lượng tài sản và năng lực quản trị rủi ro là các biến số chính. Theo methodology đã phê duyệt, tổng hợp rủi ro kinh tế vĩ mô và rủi ro ngành cho BICRA Việt Nam ở mức a-, được sử dụng làm Anchor ban đầu cho ngân hàng.')
    sections=[('3. VỊ THẾ KINH DOANH','BusinessPosition',[('TotalAssets','Tổng tài sản'),('GrossLoans','Dư nợ khách hàng'),('CustomerDeposits','Tiền gửi khách hàng')]),('4. VỐN, ĐÒN BẨY & LỢI NHUẬN','CapitalEarnings',[('CAR','Hệ số an toàn vốn (CAR)'),('ROE','ROE'),('NIM','Biên lãi ròng (NIM)')]),('5. VỊ THẾ RỦI RO','RiskPosition',[('NPL','Tỷ lệ nợ xấu (NPL)')]),('6. HUY ĐỘNG VỐN','Funding',[('CASA','CASA')]),('7. THANH KHOẢN','Liquidity',[('LDR','LDR')])]
    for title,key,charts in sections:
        doc.add_heading(title,1);doc.add_paragraph(f"Điểm: {cr['FactorScores'][key]}/"+('4' if key in {'Funding','Liquidity'} else '6')+f". {cr['FactorRationale'][key]}")
        for metric,ct in charts:doc.add_picture(_chart_metric(hist,summary,ticker,metric,ct),width=Mm(178))
    doc.add_heading('8. CẦU NỐI BICRA - SACP - ICR',1)
    bt=doc.add_table(rows=1,cols=3);bt.style='Table Grid';bt.rows[0].cells[0].text='Bước';bt.rows[0].cells[1].text='Điểm/Notch';bt.rows[0].cells[2].text='Kết quả'
    for a,b,c in [('BICRA / Anchor','Điểm xuất phát',cr['AnchorRating']),('Hồ sơ kinh doanh',f"{cr['BusinessNotch']:+d} notch",''),('Vốn & lợi nhuận',f"{cr['CapitalNotch']:+d} notch",''),('Vị thế rủi ro',f"{cr['RiskNotch']:+d} notch",''),('Huy động vốn × Thanh khoản',f"{cr['FundingLiquidityNotch']:+d} notch",''),('Điều chỉnh khác trước SACP',f"{cr['OtherInternalNotches']:+d} notch",''),('SACP','',cr['SACPRating']),('Hỗ trợ bên ngoài',f"{cr['ExternalSupportNotches']:+d} notch",''),('ICR cuối cùng','',cr['FinalRating'])]:
        cells=bt.add_row().cells;cells[0].text=a;cells[1].text=b;cells[2].text=c
    addtxt('9. HỖ TRỢ BÊN NGOÀI',f"Hỗ trợ bên ngoài được xem xét sau SACP. Điều chỉnh hiện tại: {cr['ExternalSupportNotches']:+d} notch; phải được xác lập từ bằng chứng về khả năng và mức độ sẵn sàng hỗ trợ của Chính phủ/NHNN hoặc tập đoàn sở hữu.")
    doc.add_heading('10. BẢNG ĐIỂM & PHƯƠNG PHÁP',1);t=doc.add_table(rows=1,cols=4);t.style='Table Grid';t.rows[0].cells[0].text='Yếu tố';t.rows[0].cells[1].text='Điểm';t.rows[0].cells[2].text='Notch';t.rows[0].cells[3].text='Diễn giải'
    for k,v in cr['FactorScores'].items():
        cells=t.add_row().cells;cells[0].text=FACTOR_LABELS[k];cells[1].text=f"{v}/"+('4' if k in {'Funding','Liquidity'} else '6');cells[2].text=str(cr['FactorNotches'].get(k,'—') if k not in {'Funding','Liquidity'} else '—');cells[3].text=cr['FactorRationale'][k]
    cells=t.add_row().cells;cells[0].text='Huy động vốn × Thanh khoản';cells[1].text='Ma trận Bảng 10';cells[2].text=str(cr['FundingLiquidityNotch']);cells[3].text=f"{cr['FundingDescriptor']} × {cr['LiquidityDescriptor']}"
    doc.add_heading('GIẢ ĐỊNH & HẠN CHẾ',2);doc.add_paragraph(cr['Disclaimer'])
    # Final typography normalization.
    for p in doc.paragraphs:
        if p.style.name not in ('Title','Heading 1','Heading 2') and not p._p.xpath('.//w:drawing'):
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.line_spacing=1.3;p.paragraph_format.space_before=Pt(6);p.paragraph_format.space_after=Pt(0)
        for r in p.runs:_set_lato(r,11 if p.style.name not in ('Title','Heading 1','Heading 2') else None)
    for table in doc.tables:
        for rowx in table.rows:
            for cell in rowx.cells:
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                    for r in p.runs:_set_lato(r,10)
    bio=BytesIO();doc.save(bio);bio.seek(0);return bio.getvalue()
